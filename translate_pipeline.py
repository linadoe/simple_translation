#!/usr/bin/env python3
"""
Übersetzungs-Pipeline: DOC/DOCX -> Markdown -> Englische Übersetzung (via lokalem LLM)

Konvertiert ein Word-Dokument (.doc/.docx) zu Markdown und übersetzt den Inhalt
ins Englische mithilfe eines lokal laufenden LLM. Es werden KEINE Daten ins Internet
gesendet — alle Anfragen gehen ausschließlich an den lokalen LLM-Server.
"""

import sys
import json
import re
import argparse
from pathlib import Path
from urllib.request import Request, urlopen
from urllib.error import URLError, HTTPError

from doc_to_md_robust import convert_doc_to_md


# ---------------------------------------------------------------------------
# Konfiguration
# ---------------------------------------------------------------------------
DEFAULT_LLM_URL = "http://localhost:11435"
DEFAULT_MODEL = "gpt-oss:120b"

# Maximale Zeichenanzahl pro Chunk für die Übersetzung.
# Größere Chunks = weniger API-Aufrufe, aber je nach Kontextfenster des Modells
# kann ein zu großer Chunk abgeschnitten werden.
MAX_CHUNK_CHARS = 3000

# Timeout in Sekunden für LLM-Anfragen
LLM_TIMEOUT = 300


# ---------------------------------------------------------------------------
# Hilfsfunktionen
# ---------------------------------------------------------------------------

def _validate_localhost(url: str) -> None:
    """Stellt sicher, dass die URL auf localhost zeigt — kein Internet-Zugriff."""
    from urllib.parse import urlparse
    parsed = urlparse(url)
    hostname = parsed.hostname or ""
    allowed = {"localhost", "127.0.0.1", "::1", "[::1]"}
    if hostname not in allowed:
        raise ValueError(
            f"Sicherheitsfehler: Die URL '{url}' zeigt nicht auf localhost. "
            f"Nur lokale LLM-Server sind erlaubt (hostname={hostname})."
        )


def _call_llm(prompt: str, system_prompt: str, *,
               llm_url: str, model: str) -> str:
    """
    Sendet eine Chat-Completion-Anfrage an den lokalen LLM-Server.

    Unterstützt das OpenAI-kompatible /v1/chat/completions Endpunkt-Format,
    das von Ollama, vLLM, llama.cpp-server u.a. unterstützt wird.
    """
    _validate_localhost(llm_url)

    endpoint = f"{llm_url.rstrip('/')}/api/chat"

    payload = json.dumps({
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt},
        ],
        "options": {"temperature": 0.1},
        "stream": False,
    }).encode("utf-8")

    req = Request(
        endpoint,
        data=payload,
        headers={"Content-Type": "application/json"},
        method="POST",
    )

    try:
        with urlopen(req, timeout=LLM_TIMEOUT) as resp:
            body = json.loads(resp.read().decode("utf-8"))
    except HTTPError as exc:
        raise RuntimeError(
            f"LLM-Server antwortete mit HTTP {exc.code}: {exc.read().decode('utf-8', errors='replace')}"
        ) from exc
    except URLError as exc:
        raise ConnectionError(
            f"Verbindung zum LLM-Server fehlgeschlagen ({endpoint}): {exc.reason}\n"
            "Ist der lokale LLM-Server gestartet?"
        ) from exc

    # Ollama-Antwortformat parsen
    try:
        return body["message"]["content"]
    except (KeyError, IndexError) as exc:
        raise RuntimeError(
            f"Unerwartetes Antwortformat vom LLM-Server: {json.dumps(body, indent=2)}"
        ) from exc


# ---------------------------------------------------------------------------
# Markdown-Chunking
# ---------------------------------------------------------------------------

def _split_markdown_chunks(md_text: str, max_chars: int = MAX_CHUNK_CHARS) -> list[str]:
    """
    Teilt Markdown-Text in Abschnitte auf, die das Kontextfenster des LLM
    nicht überlasten.  Versucht an Absatz- / Heading-Grenzen zu trennen.
    """
    if len(md_text) <= max_chars:
        return [md_text]

    # Aufteilen an doppelten Zeilenumbrüchen (Absatzgrenzen)
    paragraphs = re.split(r"\n{2,}", md_text)

    chunks: list[str] = []
    current_chunk = ""

    for para in paragraphs:
        candidate = (current_chunk + "\n\n" + para).strip() if current_chunk else para
        if len(candidate) <= max_chars:
            current_chunk = candidate
        else:
            if current_chunk:
                chunks.append(current_chunk)
            # Falls ein einzelner Absatz zu lang ist, nehme ihn trotzdem als eigenen Chunk
            current_chunk = para

    if current_chunk:
        chunks.append(current_chunk)

    return chunks


# ---------------------------------------------------------------------------
# Übersetzung
# ---------------------------------------------------------------------------

SYSTEM_PROMPT_MD = (
    "You are a professional translator. Translate the following Markdown text "
    "into English. Preserve ALL Markdown formatting exactly as-is: headings (#), "
    "bold (**), italic (*), links, tables, lists, code blocks, etc. "
    "Only translate the human-readable text content. "
    "Do NOT add any commentary, explanations, or notes — return ONLY the "
    "translated Markdown."
)

# Alias for backwards compatibility
SYSTEM_PROMPT = SYSTEM_PROMPT_MD

SYSTEM_PROMPT_DOCX = (
    "You are a professional translator. You will receive numbered lines of text "
    "extracted from a Word document. Translate each line into English. "
    "Return EXACTLY the same number of lines, each prefixed with the same "
    "number and pipe separator as the input (e.g. '1|translated text'). "
    "Keep empty lines as empty (e.g. '3|'). "
    "Do NOT add any commentary, explanations, or notes — return ONLY the "
    "translated lines in the exact format described."
)


def translate_markdown(md_text: str, *,
                       llm_url: str = DEFAULT_LLM_URL,
                       model: str = DEFAULT_MODEL) -> str:
    """
    Übersetzt Markdown-Text ins Englische via lokalem LLM.

    Args:
        md_text:  Der zu übersetzende Markdown-Inhalt.
        llm_url:  Basis-URL des lokalen LLM-Servers.
        model:    Modellname.

    Returns:
        Der übersetzte Markdown-Text.
    """
    chunks = _split_markdown_chunks(md_text)
    translated_parts: list[str] = []

    total = len(chunks)
    for idx, chunk in enumerate(chunks, start=1):
        print(f"  Übersetze Abschnitt {idx}/{total} ...")
        translated = _call_llm(
            prompt=chunk,
            system_prompt=SYSTEM_PROMPT,
            llm_url=llm_url,
            model=model,
        )
        translated_parts.append(translated.strip())

    return "\n\n".join(translated_parts)


# ---------------------------------------------------------------------------
# DOCX-Übersetzung
# ---------------------------------------------------------------------------

def _collect_docx_paragraphs(doc) -> list:
    """
    Sammelt alle Absätze aus dem Dokument-Body und aus Tabellenzellen.
    Gibt eine Liste von python-docx Paragraph-Objekten zurück.
    """
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)
    return paragraphs


def _chunk_paragraphs(texts: list[str], max_chars: int = MAX_CHUNK_CHARS) -> list[list[int]]:
    """
    Gruppiert Absatz-Indizes in Chunks, sodass die Gesamtzeichenzahl pro Chunk
    *max_chars* nicht überschreitet (sofern ein einzelner Absatz nicht schon
    länger ist).

    Returns:
        Liste von Listen mit Absatz-Indizes.
    """
    chunks: list[list[int]] = []
    current_chunk: list[int] = []
    current_len = 0

    for idx, text in enumerate(texts):
        # Formatierte Zeile: "idx|text\n" — Overhead einrechnen
        line_len = len(f"{idx}|{text}\n")
        if current_chunk and current_len + line_len > max_chars:
            chunks.append(current_chunk)
            current_chunk = []
            current_len = 0
        current_chunk.append(idx)
        current_len += line_len

    if current_chunk:
        chunks.append(current_chunk)

    return chunks


def _translate_paragraph_chunk(
    indices: list[int],
    texts: list[str],
    *,
    llm_url: str,
    model: str,
) -> dict[int, str]:
    """
    Übersetzt einen Chunk von Absätzen via LLM.
    Gibt ein dict {index: übersetzter_text} zurück.
    """
    # Nummeriertes Format aufbauen
    lines = [f"{i}|{texts[i]}" for i in indices]
    prompt = "\n".join(lines)

    raw = _call_llm(
        prompt=prompt,
        system_prompt=SYSTEM_PROMPT_DOCX,
        llm_url=llm_url,
        model=model,
    )

    # Antwort parsen: "idx|übersetzter Text"
    result: dict[int, str] = {}
    for line in raw.strip().splitlines():
        line = line.strip()
        if not line:
            continue
        if "|" in line:
            key, _, value = line.partition("|")
            try:
                result[int(key.strip())] = value
            except ValueError:
                continue

    return result


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """
    Ersetzt den Text eines Absatzes unter Beibehaltung der Formatierung.

    - Ein Run: Text direkt ersetzen (Format bleibt komplett erhalten).
    - Mehrere Runs: Übersetzung in den ersten Run, restliche leeren.
    """
    runs = paragraph.runs
    if not runs:
        return
    if len(runs) == 1:
        runs[0].text = new_text
    else:
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""


def translate_docx(
    input_file: str,
    output_file: str | None = None,
    *,
    llm_url: str = DEFAULT_LLM_URL,
    model: str = DEFAULT_MODEL,
) -> str:
    """
    Übersetzt eine .docx-Datei ins Englische und speichert das Ergebnis
    als neue .docx-Datei. Word-Formatierung (Styles, Schriftarten, Tabellen)
    wird beibehalten.

    Args:
        input_file:   Pfad zur Eingabedatei (.docx).
        output_file:  Pfad für die übersetzte .docx-Datei.
                      Standard: <input>_en.docx
        llm_url:      Basis-URL des lokalen LLM-Servers.
        model:        Modellname des LLM.

    Returns:
        Pfad zur erzeugten Ausgabedatei.
    """
    from docx import Document

    _validate_localhost(llm_url)

    input_path = Path(input_file)
    if not input_path.exists():
        raise FileNotFoundError(f"Datei nicht gefunden: {input_path}")

    if input_path.suffix.lower() != ".docx":
        raise ValueError(
            f"Nicht unterstütztes Dateiformat: '{input_path.suffix}'. "
            "translate_docx erwartet eine .docx-Datei."
        )

    # --- Dokument öffnen ---
    print("=" * 60)
    print("DOCX-Übersetzung: Dokument wird geladen")
    print("=" * 60)
    doc = Document(str(input_path))

    # --- Alle Absätze sammeln ---
    all_paragraphs = _collect_docx_paragraphs(doc)
    texts = [p.text for p in all_paragraphs]

    # Nur nicht-leere Absätze übersetzen
    non_empty_indices = [i for i, t in enumerate(texts) if t.strip()]
    print(f"  {len(all_paragraphs)} Absätze gefunden, "
          f"davon {len(non_empty_indices)} nicht-leer.")

    if not non_empty_indices:
        print("  Keine übersetzbaren Absätze gefunden.")
    else:
        # --- Chunking ---
        non_empty_texts = [texts[i] for i in non_empty_indices]
        # _chunk_paragraphs arbeitet auf lokalen Indizes,
        # wir müssen sie auf die non_empty_indices abbilden.
        local_chunks = _chunk_paragraphs(non_empty_texts)
        total_chunks = len(local_chunks)

        print(f"  Aufgeteilt in {total_chunks} Chunk(s).\n")
        print("=" * 60)
        print("Übersetzung via lokalem LLM")
        print(f"  Server: {llm_url}")
        print(f"  Modell: {model}")
        print("=" * 60)

        # --- Übersetzen ---
        translated_map: dict[int, str] = {}  # global index -> übersetzter Text
        for chunk_idx, local_indices in enumerate(local_chunks, start=1):
            print(f"  Übersetze Chunk {chunk_idx}/{total_chunks} "
                  f"({len(local_indices)} Absätze) ...")

            # Globale Indizes für diesen Chunk
            global_indices = [non_empty_indices[li] for li in local_indices]
            result = _translate_paragraph_chunk(
                global_indices, texts,
                llm_url=llm_url, model=model,
            )
            translated_map.update(result)

        # --- Übersetzungen zurückschreiben ---
        replaced = 0
        for idx in non_empty_indices:
            if idx in translated_map:
                _replace_paragraph_text(all_paragraphs[idx], translated_map[idx])
                replaced += 1

        print(f"\n  {replaced}/{len(non_empty_indices)} Absätze übersetzt "
              f"und zurückgeschrieben.")

    # --- Speichern ---
    if output_file is None:
        output_path = input_path.with_name(input_path.stem + "_en.docx")
    else:
        output_path = Path(output_file)

    doc.save(str(output_path))

    print("\n" + "=" * 60)
    print(f"Übersetzung abgeschlossen: {output_path}")
    print("=" * 60)

    return str(output_path)


# ---------------------------------------------------------------------------
# Haupt-Pipeline (Markdown)
# ---------------------------------------------------------------------------

def translate_document(
    input_file: str,
    output_file: str | None = None,
    *,
    llm_url: str = DEFAULT_LLM_URL,
    model: str = DEFAULT_MODEL,
    conversion_method: str = "auto",
) -> str:
    """
    Komplette Pipeline: DOC/DOCX → Markdown → Englische Übersetzung.

    Args:
        input_file:         Pfad zur Eingabedatei (.doc / .docx).
        output_file:        Pfad für die übersetzte .md-Datei.
                            Standard: <input>_en.md
        llm_url:            Basis-URL des lokalen LLM-Servers.
        model:              Modellname des LLM.
        conversion_method:  Konvertierungsmethode für doc→md
                            ('auto', 'mammoth', 'pypandoc', 'python-docx', 'zipfile').

    Returns:
        Den übersetzten Markdown-Inhalt als String.
    """
    _validate_localhost(llm_url)

    input_path = Path(input_file)
    if not input_path.exists():
        raise FileNotFoundError(f"Datei nicht gefunden: {input_path}")

    suffix = input_path.suffix.lower()
    if suffix not in {".doc", ".docx"}:
        raise ValueError(
            f"Nicht unterstütztes Dateiformat: '{suffix}'. "
            "Bitte eine .doc oder .docx Datei angeben."
        )

    # --- Schritt 1: Konvertierung DOC/DOCX → Markdown ---
    print("=" * 60)
    print("SCHRITT 1: Konvertierung DOC/DOCX -> Markdown")
    print("=" * 60)

    # Temporäre Markdown-Datei für die Zwischenstufe
    intermediate_md = input_path.with_suffix(".md")
    md_content = convert_doc_to_md(
        input_path, intermediate_md, method=conversion_method
    )

    print(f"\nMarkdown erstellt: {intermediate_md}")
    print(f"Inhaltslänge: {len(md_content)} Zeichen\n")

    # --- Schritt 2: Übersetzung via lokalem LLM ---
    print("=" * 60)
    print("SCHRITT 2: Übersetzung ins Englische (lokaler LLM)")
    print(f"  Server: {llm_url}")
    print(f"  Modell: {model}")
    print("=" * 60)

    translated = translate_markdown(md_content, llm_url=llm_url, model=model)

    # --- Schritt 3: Ergebnis speichern ---
    if output_file is None:
        output_path = input_path.with_name(input_path.stem + "_en.md")
    else:
        output_path = Path(output_file)

    output_path.write_text(translated, encoding="utf-8")

    print("\n" + "=" * 60)
    print(f"Übersetzung abgeschlossen: {output_path}")
    print(f"Ergebnislänge: {len(translated)} Zeichen")
    print("=" * 60)

    return translated


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="DOC/DOCX → Markdown → Englische Übersetzung (lokaler LLM)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Beispiele:\n"
            "  python translate_pipeline.py dokument.docx\n"
            "  python translate_pipeline.py dokument.docx -o ergebnis_en.md\n"
            "  python translate_pipeline.py dokument.docx --model gpt-oss:120b\n"
            "  python translate_pipeline.py dokument.docx --method mammoth\n"
        ),
    )
    parser.add_argument(
        "input_file",
        help="Eingabedatei (.doc oder .docx)",
    )
    parser.add_argument(
        "-o", "--output",
        default=None,
        help="Ausgabedatei für die übersetzte Markdown-Datei (Standard: <input>_en.md)",
    )
    parser.add_argument(
        "--llm-url",
        default=DEFAULT_LLM_URL,
        help=f"URL des lokalen LLM-Servers (Standard: {DEFAULT_LLM_URL})",
    )
    parser.add_argument(
        "--model",
        default=DEFAULT_MODEL,
        help=f"LLM-Modellname (Standard: {DEFAULT_MODEL})",
    )
    parser.add_argument(
        "--method",
        default="auto",
        choices=["auto", "mammoth", "pypandoc", "python-docx", "zipfile"],
        help="Konvertierungsmethode für DOC/DOCX → Markdown (Standard: auto)",
    )

    args = parser.parse_args()

    try:
        translate_document(
            input_file=args.input_file,
            output_file=args.output,
            llm_url=args.llm_url,
            model=args.model,
            conversion_method=args.method,
        )
    except (FileNotFoundError, ValueError, ConnectionError, RuntimeError) as exc:
        print(f"\nFehler: {exc}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
