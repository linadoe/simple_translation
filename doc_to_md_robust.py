#!/usr/bin/env python3
"""
DOC/DOCX to Markdown Converter - Robuste Version
Mehrere Methoden für maximale Kompatibilität
"""

import sys
import subprocess
from pathlib import Path


def check_and_install_dependencies():
    """Prüft und installiert fehlende Pakete"""
    packages = {
        'mammoth': 'mammoth',
        'python-docx': 'docx',
        'pypandoc': 'pypandoc'
    }
    
    missing = []
    for pip_name, import_name in packages.items():
        try:
            __import__(import_name)
        except ImportError:
            missing.append(pip_name)
    
    if missing:
        print(f"Fehlende Pakete: {', '.join(missing)}")
        print("\nInstalliere mit:")
        print(f"pip install {' '.join(missing)} --break-system-packages")
        return False
    return True


def method_mammoth(docx_path, output_path):
    """
    Methode 1: mammoth (empfohlen)
    - Beste Formatierung
    - Unterstützt fett, kursiv, Listen
    """
    try:
        import mammoth
    except ImportError:
        print("✗ mammoth nicht installiert")
        return None
    
    print("Verwende mammoth...")
    with open(docx_path, "rb") as docx_file:
        result = mammoth.convert_to_markdown(docx_file)
        markdown = result.value
        
        if result.messages:
            print("Hinweise:")
            for msg in result.messages:
                print(f"  {msg}")
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(markdown)
    
    return markdown


def method_pypandoc(docx_path, output_path):
    """
    Methode 2: pypandoc
    - Nutzt Pandoc (muss separat installiert sein)
    - Sehr gute Konvertierung
    """
    try:
        import pypandoc
    except ImportError:
        print("✗ pypandoc nicht installiert")
        return None
    
    print("Verwende pypandoc...")
    try:
        markdown = pypandoc.convert_file(str(docx_path), 'md')
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown)
        return markdown
    except RuntimeError as e:
        print(f"✗ Pandoc nicht gefunden: {e}")
        print("Installiere Pandoc: https://pandoc.org/installing.html")
        return None


def method_python_docx(docx_path, output_path):
    """
    Methode 3: python-docx
    - Fallback-Methode
    - Grundlegende Konvertierung
    """
    try:
        from docx import Document
    except ImportError:
        print("✗ python-docx Import fehlgeschlagen")
        print("\nProblem beheben:")
        print("1. Deinstalliere alte Pakete:")
        print("   pip uninstall docx python-docx")
        print("2. Installiere python-docx neu:")
        print("   pip install python-docx --break-system-packages")
        return None
    
    print("Verwende python-docx...")
    doc = Document(docx_path)
    lines = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        
        # Überschriften
        if para.style.name.startswith('Heading'):
            try:
                level = int(para.style.name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            lines.append(f"{'#' * level} {text}")
        else:
            lines.append(text)
    
    # Tabellen
    for table in doc.tables:
        lines.append("")
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
        lines.append("")
    
    markdown = "\n".join(lines)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(markdown)
    
    return markdown


def method_zipfile(docx_path, output_path):
    """
    Methode 4: Manuelle Extraktion (letzter Ausweg)
    - Extrahiert nur Text ohne Formatierung
    - Funktioniert ohne Dependencies
    """
    import zipfile
    import xml.etree.ElementTree as ET
    
    print("Verwende manuelle Extraktion (nur Text)...")
    
    try:
        with zipfile.ZipFile(docx_path) as docx:
            xml_content = docx.read('word/document.xml')
            tree = ET.fromstring(xml_content)
            
            # Namespace für Word XML
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            paragraphs = []
            for para in tree.findall('.//w:p', ns):
                texts = para.findall('.//w:t', ns)
                para_text = ''.join(t.text for t in texts if t.text)
                if para_text.strip():
                    paragraphs.append(para_text)
            
            markdown = '\n\n'.join(paragraphs)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown)
            
            print("⚠ Warnung: Nur Text extrahiert, keine Formatierung")
            return markdown
    
    except Exception as e:
        print(f"✗ Manuelle Extraktion fehlgeschlagen: {e}")
        return None


def convert_doc_to_md(input_path, output_path=None, method='auto'):
    """
    Hauptfunktion - versucht verschiedene Methoden
    
    Args:
        input_path: Pfad zur .docx-Datei
        output_path: Pfad für .md-Datei (optional)
        method: 'auto', 'mammoth', 'pypandoc', 'python-docx', 'zipfile'
    """
    input_path = Path(input_path)
    
    if not input_path.exists():
        raise FileNotFoundError(f"Datei nicht gefunden: {input_path}")
    
    if output_path is None:
        output_path = input_path.with_suffix('.md')
    else:
        output_path = Path(output_path)
    
    print(f"Konvertiere: {input_path}")
    print(f"Ausgabe: {output_path}\n")
    
    methods_map = {
        'mammoth': method_mammoth,
        'pypandoc': method_pypandoc,
        'python-docx': method_python_docx,
        'zipfile': method_zipfile
    }
    
    if method != 'auto':
        if method not in methods_map:
            raise ValueError(f"Unbekannte Methode: {method}")
        result = methods_map[method](input_path, output_path)
        if result is None:
            raise RuntimeError(f"Methode {method} fehlgeschlagen")
        return result
    
    # Auto-Modus: Versuche alle Methoden
    for method_name, method_func in methods_map.items():
        print(f"\n--- Versuche {method_name} ---")
        result = method_func(input_path, output_path)
        if result is not None:
            print(f"\n✓ Erfolgreich mit {method_name}!")
            return result
    
    raise RuntimeError("Alle Konvertierungsmethoden fehlgeschlagen")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("DOC/DOCX zu Markdown Konverter")
        print("=" * 50)
        print("\nVerwendung:")
        print("  python doc_to_md_robust.py <input.docx> [output.md] [methode]")
        print("\nBeispiele:")
        print("  python doc_to_md_robust.py document.docx")
        print("  python doc_to_md_robust.py document.docx output.md")
        print("  python doc_to_md_robust.py document.docx output.md mammoth")
        print("\nVerfügbare Methoden:")
        print("  auto         - Automatisch beste Methode wählen (Standard)")
        print("  mammoth      - Beste Formatierung (empfohlen)")
        print("  pypandoc     - Benötigt Pandoc")
        print("  python-docx  - Grundlegende Konvertierung")
        print("  zipfile      - Nur Text, keine Dependencies")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    method = sys.argv[3] if len(sys.argv) > 3 else 'auto'
    
    try:
        convert_doc_to_md(input_file, output_file, method)
        print("\n" + "=" * 50)
        print("✓ Konvertierung erfolgreich abgeschlossen!")
    except Exception as e:
        print(f"\n✗ Fehler: {e}")
        sys.exit(1)


